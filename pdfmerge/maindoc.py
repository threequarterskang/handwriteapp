import sqlite3
import json
from pathlib import Path
import glob
from docx import Document
from docx.shared import Pt, RGBColor
import re
import os
import shutil

QUERY_DB_NAME = "placeholder.db"
DATA_DB_NAME = "total.db"
PDF_DIRECTORY = "./pdf/"
folder = Path("./config")
template_path = "./docx/"
output_path = "./docxout/"

# 定义字体和颜色
FONT_NAME = "TekitouPoem"  # 你指定的字体
FONT_COLOR = RGBColor(0, 0, 0)  # 黑色
FONT_SIZE = Pt(15)  # 默认字体大小

def ensure_dir_exists(dir_path):
    """确保目录存在"""
    if not os.path.exists(dir_path):
        os.makedirs(dir_path, exist_ok=True)
        print(f"创建目录: {dir_path}")

def debug_table_structure(doc, template_name):
    """专门调试表格结构，特别是表头"""
    print(f"\n{'='*60}")
    print(f"详细调试表格结构: {template_name}")
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n表格 {table_idx}: {len(table.rows)} 行, {len(table.columns)} 列")
        
        # 检查每一行
        for row_idx, row in enumerate(table.rows):
            print(f"  行 {row_idx}:")
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    print(f"    列 {cell_idx}: '{cell_text}'")
                    
                    # 检查段落和运行
                    for para_idx, para in enumerate(cell.paragraphs):
                        if para.text.strip():
                            print(f"      段落 {para_idx}: '{para.text}'")
                            for run_idx, run in enumerate(para.runs):
                                if run.text.strip():
                                    print(f"        运行 {run_idx}: '{run.text}'")
                                    # 打印运行属性
                                    print(f"          字体: {run.font.name}, 加粗: {run.font.bold}, 斜体: {run.font.italic}")

def find_placeholders_in_table_header(doc):
    """专门查找表格标题行中的占位符"""
    placeholders = set()
    
    for table in doc.tables:
        if len(table.rows) > 0:
            header_row = table.rows[0]  # 假设第一行是标题行
            
            for cell in header_row.cells:
                for para in cell.paragraphs:
                    # 查找 {数字} 格式
                    matches = re.findall(r'\{(\d+)\}', para.text)
                    for match in matches:
                        placeholders.add(f"{{{match}}}")
                    
                    # 查找其他可能的格式
                    patterns = [
                        r'\[(\d+)\]',      # [1]
                        r'<(\d+)>',        # <1>
                        r'\((\d+)\)',      # (1)
                        r'【(\d+)】',      # 【1】
                        r'（(\d+)）',      # （1）
                        r'〈(\d+)〉',      # 〈1〉
                    ]
                    
                    for pattern in patterns:
                        matches = re.findall(pattern, para.text)
                        for match in matches:
                            placeholders.add(f"{{{match}}}")
    
    return sorted(list(placeholders), key=lambda x: int(re.findall(r'\d+', x)[0]) if re.findall(r'\d+', x) else 0)

def replace_in_table_header(table, replacements, table_idx):
    """专门替换表格标题行中的占位符"""
    replaced_count = 0
    
    if len(table.rows) > 0:
        header_row = table.rows[0]
        
        for cell_idx, cell in enumerate(header_row.cells):
            for para_idx, para in enumerate(cell.paragraphs):
                original_text = para.text
                
                for placeholder, value in replacements.items():
                    if placeholder in original_text:
                        # 保存格式信息
                        old_runs = list(para.runs)
                        is_bold = False
                        is_italic = False
                        alignment = para.alignment
                        
                        if old_runs:
                            is_bold = old_runs[0].font.bold
                            is_italic = old_runs[0].font.italic
                        
                        # 清空段落
                        para.clear()
                        
                        # 替换文本
                        new_text = original_text.replace(placeholder, str(value))
                        new_run = para.add_run(new_text)
                        
                        # 应用字体
                        new_run.font.name = FONT_NAME
                        new_run.font.color.rgb = FONT_COLOR
                        new_run.font.size = FONT_SIZE
                        
                        # 恢复格式
                        if is_bold:
                            new_run.font.bold = True
                        if is_italic:
                            new_run.font.italic = True
                        
                        # 恢复对齐
                        para.alignment = alignment
                        
                        replaced_count += 1
                        print(f"表格{table_idx}标题行 列{cell_idx}: 替换 '{placeholder}' -> '{value}'")
                        break
    
    return replaced_count

def replace_docx_placeholders(template_name, output_name, replacements):
    """增强版：替换DOCX中的占位符，特别处理表格标题行"""
    if not os.path.exists(template_name):
        print(f"错误: 模板文件不存在 {template_name}")
        return False
    
    try:
        doc = Document(template_name)
        
        if not replacements:
            print("警告: 没有替换内容")
            doc.save(output_name)
            return True
        
        print(f"开始处理: {template_name}")
        print(f"替换映射: {replacements}")
        
        # 1. 首先专门查找表格标题行中的占位符
        header_placeholders = find_placeholders_in_table_header(doc)
        print(f"表格标题行中找到的占位符: {header_placeholders}")
        
        # 如果没有找到，显示详细的表格结构
        if not header_placeholders:
            print("警告: 未在表格标题行中找到占位符，显示详细结构...")
            debug_table_structure(doc, template_name)
        
        total_replaced = 0
        
        # 2. 专门处理表格标题行
        for table_idx, table in enumerate(doc.tables):
            header_replaced = replace_in_table_header(table, replacements, table_idx)
            total_replaced += header_replaced
        
        # 3. 处理表格其他行
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows[1:] if len(table.rows) > 1 else []):  # 跳过标题行
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        original_text = para.text
                        
                        for placeholder, value in replacements.items():
                            if placeholder in original_text:
                                # 保存格式
                                old_runs = list(para.runs)
                                
                                # 清空并替换
                                para.clear()
                                new_text = original_text.replace(placeholder, str(value))
                                new_run = para.add_run(new_text)
                                
                                # 应用字体
                                new_run.font.name = FONT_NAME
                                new_run.font.color.rgb = FONT_COLOR
                                new_run.font.size = FONT_SIZE
                                
                                total_replaced += 1
                                print(f"表格{table_idx}行{row_idx+1}列{cell_idx}: 替换 '{placeholder}' -> '{value}'")
                                break
        
        # 4. 处理段落
        for para_idx, para in enumerate(doc.paragraphs):
            original_text = para.text
            
            for placeholder, value in replacements.items():
                if placeholder in original_text:
                    # 保存对齐方式
                    alignment = para.alignment
                    
                    # 清空并替换
                    para.clear()
                    new_text = original_text.replace(placeholder, str(value))
                    new_run = para.add_run(new_text)
                    
                    # 应用字体
                    new_run.font.name = FONT_NAME
                    new_run.font.color.rgb = FONT_COLOR
                    new_run.font.size = FONT_SIZE
                    
                    # 恢复对齐
                    para.alignment = alignment
                    
                    total_replaced += 1
                    print(f"段落{para_idx}: 替换 '{placeholder}' -> '{value}'")
                    break
        
        # 5. 处理页眉页脚
        for section in doc.sections:
            # 页眉
            for header in [section.header, section.first_page_header]:
                for para in header.paragraphs:
                    original_text = para.text
                    for placeholder, value in replacements.items():
                        if placeholder in original_text:
                            para.clear()
                            new_run = para.add_run(original_text.replace(placeholder, str(value)))
                            new_run.font.name = FONT_NAME
                            new_run.font.color.rgb = FONT_COLOR
                            new_run.font.size = FONT_SIZE
                            total_replaced += 1
                            break
            
            # 页脚
            for footer in [section.footer, section.first_page_footer]:
                for para in footer.paragraphs:
                    original_text = para.text
                    for placeholder, value in replacements.items():
                        if placeholder in original_text:
                            para.clear()
                            new_run = para.add_run(original_text.replace(placeholder, str(value)))
                            new_run.font.name = FONT_NAME
                            new_run.font.color.rgb = FONT_COLOR
                            new_run.font.size = FONT_SIZE
                            total_replaced += 1
                            break
        
        if total_replaced > 0:
            print(f"✓ 成功替换了 {total_replaced} 处占位符")
            doc.save(output_name)
            return True
        else:
            print("✗ 警告: 未找到任何占位符进行替换")
            
            # 尝试更深入的搜索
            print("\n尝试深度搜索占位符...")
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            all_text.append(cell.text)
            
            print(f"文档中的所有文本片段 ({len(all_text)} 个):")
            for i, text in enumerate(all_text[:20]):  # 只显示前20个
                print(f"  {i}: '{text}'")
            
            # 保存未修改的文档
            doc.save(output_name)
            return False
            
    except Exception as e:
        print(f"✗ 处理失败 {template_name} -> {output_name}: {e}")
        import traceback
        traceback.print_exc()
        return False

def find_template_file(tbn, template_dir):
    """查找模板文件"""
    # 尝试多种可能的文件名格式
    patterns = [
        f"* {tbn}.docx",  # 匹配 "表名.docx" 或 "* 表名.docx"
        f"{tbn}.docx",    # 直接匹配表名
        f"*{tbn}*.docx",  # 模糊匹配
    ]
    
    for pattern in patterns:
        search_path = os.path.join(template_dir, pattern)
        files = glob.glob(search_path, recursive=False)
        if files:
            print(f"找到模板文件: {files[0]} (模式: {pattern})")
            return files[0]
    
    # 如果没有找到，列出所有可用的模板文件
    print(f"在 {template_dir} 中找到的模板文件:")
    all_files = glob.glob(os.path.join(template_dir, "*.docx"))
    for file in all_files:
        print(f"  - {os.path.basename(file)}")
    
    return None

def inspect_template_deep(tbn):
    """深度检查模板文件内容，特别是表格标题行"""
    template_file = find_template_file(tbn, template_path)
    if not template_file:
        print(f"未找到模板文件: {tbn}")
        return
    
    print(f"\n{'='*60}")
    print(f"深度检查模板文件: {template_file}")
    
    try:
        doc = Document(template_file)
        
        # 1. 检查表格标题行
        print(f"\n文档中的表格数量: {len(doc.tables)}")
        for t, table in enumerate(doc.tables):
            print(f"\n表格 {t}: {len(table.rows)} 行, {len(table.columns)} 列")
            
            if len(table.rows) > 0:
                print(f"标题行 (第0行) 内容:")
                header_row = table.rows[0]
                for c, cell in enumerate(header_row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"  列{c}: '{cell_text}'")
                        
                        # 检查是否有占位符
                        matches = re.findall(r'\{(\d+)\}', cell_text)
                        if matches:
                            print(f"    包含占位符: {[f'{{{m}}}' for m in matches]}")
                        
                        # 检查段落
                        for p, para in enumerate(cell.paragraphs):
                            if para.text.strip():
                                print(f"    段落{p}: '{para.text}'")
        
        # 2. 查找所有可能的占位符
        print(f"\n{'='*60}")
        print("搜索所有占位符:")
        
        all_placeholders = set()
        
        # 在段落中搜索
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                matches = re.findall(r'\{(\d+)\}', para.text)
                for match in matches:
                    all_placeholders.add(f"{{{match}}}")
                if matches:
                    print(f"段落{i}: '{para.text[:50]}...'")
                    print(f"  找到占位符: {[f'{{{m}}}' for m in matches]}")
        
        # 在表格中搜索
        for t, table in enumerate(doc.tables):
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    for p, para in enumerate(cell.paragraphs):
                        if para.text.strip():
                            matches = re.findall(r'\{(\d+)\}', para.text)
                            for match in matches:
                                all_placeholders.add(f"{{{match}}}")
                            if matches:
                                print(f"表格{t}行{r}列{c}: '{para.text[:50]}...'")
                                print(f"  找到占位符: {[f'{{{m}}}' for m in matches]}")
        
        if all_placeholders:
            print(f"\n文档中所有占位符: {sorted(list(all_placeholders), key=lambda x: int(x.strip('{}')))}")
        else:
            print("警告: 未找到任何占位符!")
        
    except Exception as e:
        print(f"检查模板失败: {e}")

def main():
    # 确保目录存在
    ensure_dir_exists(template_path)
    ensure_dir_exists(output_path)
    
    # 检查数据库文件是否存在
    if not os.path.exists(DATA_DB_NAME):
        print(f"错误: 数据数据库文件不存在 {DATA_DB_NAME}")
        return
    
    if not os.path.exists(QUERY_DB_NAME):
        print(f"错误: 查询数据库文件不存在 {QUERY_DB_NAME}")
        return
    
    # 连接数据库
    try:
        conn1 = sqlite3.connect(DATA_DB_NAME)
        conn2 = sqlite3.connect(QUERY_DB_NAME)
        conn1.row_factory = sqlite3.Row
        
        cur1 = conn1.cursor()
        cur2 = conn2.cursor()

        # 获取字段规则
        res = cur2.execute(
            'SELECT "英文名字", "数据最大范围" FROM field_rule'
        )
        
        if res is None:
            raise ValueError("无法获取字段规则")
            
        mapping = dict(res.fetchall())
        print(f"加载了 {len(mapping)} 个字段规则")

        # 获取所有表格
        cur1.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'"
        )
        
        tables = [row[0] for row in cur1.fetchall()]

        if not tables:
            print("警告: 数据库中没有表格")
            return
        
        print(f"找到 {len(tables)} 个表格: {tables}")
        
        processed_count = 0
        
        # 处理每个表格
        for tbn in tables:
            if tbn in mapping:
                print(f"\n{'='*60}")
                print(f"处理表格: {tbn}")
                
                # 深度检查模板
                inspect_template_deep(tbn)
                
                # 查找模板文件
                template_file = find_template_file(tbn, template_path)
                if not template_file:
                    print(f"警告: 未找到表格 {tbn} 的模板文件")
                    continue
                
                print(f"使用模板: {template_file}")
                
                # 查询数据
                cur1.execute(f'SELECT * FROM "{tbn}"')
                rows = cur1.fetchall()
                
                if not rows:
                    print(f"表格 {tbn} 中没有数据")
                    continue
                
                max_column = int(mapping[tbn])
                row_count = len(rows)
                
                print(f"表格 {tbn} 有 {row_count} 行数据，最大列数: {max_column}")
                
                # 处理每一行
                for row_index, row in enumerate(rows):
                    placeholer_map = {}
                    codename = None
                    
                    print(f"\n处理第 {row_index + 1} 行数据:")
                    
                    # 提取数据
                    for idx in range(1, max_column + 1):
                        column_name = f"{{{idx}}}"
                        if column_name in row.keys():
                            text = row[column_name]
                            
                            if text is None:
                                placeholer_map[column_name] = ""
                                print(f"  {column_name}: [空值]")
                            else:
                                placeholer_map[column_name] = str(text)
                                print(f"  {column_name}: {text}")
                            
                            # 保存第7列作为文件名
                            if idx == 7 and text is not None:
                                codename = str(text)
                    
                    # 如果没有codename，使用其他标识
                    if not codename:
                        # 尝试使用第一列或创建组合名称
                        first_col = f"{{{1}}}"
                        if first_col in row.keys() and row[first_col]:
                            codename = f"{tbn}_{row[first_col]}"
                        else:
                            codename = f"{tbn}_{row_index + 1}"
                    
                    # 清理文件名
                    import re
                    codename = re.sub(r'[<>:"/\\|?*]', '_', codename)
                    codename = codename[:100]  # 限制文件名长度
                    
                    # 生成输出文件名
                    output_name = os.path.join(output_path, f"{codename}.docx")
                    
                    # 如果文件已存在，添加序号
                    counter = 1
                    while os.path.exists(output_name):
                        output_name = os.path.join(output_path, f"{codename}_{counter}.docx")
                        counter += 1
                    
                    # 替换并生成文档
                    if placeholer_map:
                        print(f"\n生成文档: {output_name}")
                        print(f"应用字体: {FONT_NAME}, 颜色: 黑色, 大小: {FONT_SIZE.pt}pt")
                        
                        if replace_docx_placeholders(template_file, output_name, placeholer_map):
                            processed_count += 1
                        else:
                            print(f"处理失败: 行 {row_index + 1}")
                    else:
                        print(f"跳过: 行 {row_index + 1} 没有数据")
            
            else:
                print(f"跳过表格 {tbn}: 不在字段规则中")
        
        print(f"\n{'='*60}")
        print(f"处理完成! 共生成 {processed_count} 个文档")
        print(f"输出目录: {output_path}")
        
        # 显示生成的文档
        if os.path.exists(output_path):
            output_files = os.listdir(output_path)
            if output_files:
                print(f"生成的文档: {output_files[:10]}")  # 显示前10个
                if len(output_files) > 10:
                    print(f"... 共 {len(output_files)} 个文件")
                    
    except sqlite3.Error as e:
        print(f"数据库错误: {e}")
    except Exception as e:
        print(f"程序错误: {e}")
        import traceback
        traceback.print_exc()
    finally:
        # 确保关闭数据库连接
        if 'conn1' in locals() and conn1:
            conn1.close()
        if 'conn2' in locals() and conn2:
            conn2.close()

# 专门的表格标题行调试工具
def debug_table_header_only():
    """专门调试表格标题行"""
    template_name = input("输入模板文件名或表名: ").strip()
    
    template_file = find_template_file(template_name, template_path)
    if not template_file:
        print(f"未找到模板文件: {template_name}")
        return
    
    print(f"\n{'='*60}")
    print(f"专门调试表格标题行: {template_file}")
    
    try:
        doc = Document(template_file)
        debug_table_structure(doc, template_file)
        
        # 专门检查表格标题行
        print(f"\n{'='*60}")
        print("表格标题行详细分析:")
        
        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) > 0:
                print(f"\n表格 {table_idx} 标题行:")
                header_row = table.rows[0]
                
                for cell_idx, cell in enumerate(header_row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"  列{cell_idx}: '{cell_text}'")
                        
                        # 逐个字符分析
                        print(f"    字符分析: {[char for char in cell_text]}")
                        print(f"    ASCII码: {[ord(char) for char in cell_text]}")
                        
                        # 检查是否是特殊字符
                        for char in cell_text:
                            if char == '{' or char == '}' or char == '[' or char == ']' or char == '(' or char == ')':
                                print(f"    特殊字符: '{char}' (ASCII: {ord(char)})")
                        
                        # 检查是否有占位符模式
                        patterns = [r'\{(\d+)\}', r'\[(\d+)\]', r'<(\d+)>', r'\((\d+)\)']
                        for pattern in patterns:
                            matches = re.findall(pattern, cell_text)
                            if matches:
                                print(f"    匹配模式 '{pattern}': {matches}")
            
    except Exception as e:
        print(f"调试失败: {e}")

if __name__ == "__main__":
    print("1. 运行主程序")
    print("2. 检查模板文件")
    print("3. 专门调试表格标题行")
    choice = input("请选择 (1/2/3): ").strip()
    
    if choice == "2":
        # 使用之前定义的检查函数
        template_name = input("输入模板文件名或表名: ").strip()
        template_file = find_template_file(template_name, template_path)
        if template_file:
            try:
                doc = Document(template_file)
                debug_table_structure(doc, template_file)
            except Exception as e:
                print(f"检查失败: {e}")
    elif choice == "3":
        debug_table_header_only()
    else:
        main()
    
    print("\n完成!")